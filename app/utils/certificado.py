# Documentacion de la librería docx - https://python-docx.readthedocs.io/
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from docx.shared import Pt
from docx.shared import Inches
from io import BytesIO
from app.schemas.estudiante import Estudiante
import os


def generar_certificado(id_estudiante: int, path: str) -> BytesIO:
    """
    Generar un documento en en word que contiene todos los datos del estudiante y las actividades realizadas.
    La función acepta como parámetros el id del estudiante y la ruta (root) de la carpeta de la aplicación
    Esta función devuelve un objeto BytesIO el cual contiene el documento creado.
    """
    # Funcion para poner el mes
    estudiante = Estudiante.query.get(id_estudiante)
    coordinador = "________________________________________\nLcda. Rocío Campoverde Carpio, Mg.\nCoordinadora General \n CDR “Juan Bautista Vázquez”"

    def mes(num):
        meses = [
            "enero",
            "febrero",
            "marzo",
            "abril",
            "mayo",
            "junio",
            "julio",
            "agosto",
            "septiembre",
            "octubre",
            "noviembre",
            "diciembre",
        ]
        return meses[num - 1]

    # Datos
    fecha_actual = datetime.datetime.now()
    fecha = f"{fecha_actual.day} de {mes(fecha_actual.month)} de {fecha_actual.year}"

    # Creación de la instancia de clase Document
    document = Document()
    certif_estud = BytesIO()

    # Encabezado
    ruta_imagen = os.path.join(path, "static/img", "ucuenca.jpg")
    header = document.sections[0].header
    encabezado = header.paragraphs[0]
    logo_UCUENCA = encabezado.add_run()
    logo_UCUENCA.add_picture(ruta_imagen, width=Inches(2))

    encabezado = document.add_paragraph()
    texto_encabezado = encabezado.add_run("CERTIFICA\n")
    texto_encabezado.bold = True
    fuente_texto_encabezado = texto_encabezado.font
    fuente_texto_encabezado.size = Pt(18)
    fuente_texto_encabezado.name = "Arial"
    formato_encabezado = encabezado.paragraph_format
    formato_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Texto central
    parrafo = document.add_paragraph()
    parrafo1 = parrafo.add_run("Que ")
    parrafo2 = parrafo.add_run(f"{estudiante.nombre} ")
    parrafo3 = parrafo.add_run(" con cédula No. ")
    parrafo4 = parrafo.add_run(f"{estudiante.cedula}, ")
    parrafo5 = parrafo.add_run(
        f" estudiante de la {estudiante.facultad}, {estudiante.carrera}, ha cumplido con el trabajo de {estudiante.total_horas} HORAS de servicio académico en esta dependencia. Sus labores de servicio estuvieron deidcadas a:                                                          \n"
    )

    formato_parrafo = parrafo.paragraph_format
    formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    fuente_parrafo1 = parrafo1.font
    fuente_parrafo1.size = Pt(12)
    fuente_parrafo1.name = "Arial"
    parrafo2.bold = True
    fuente_parrafo2 = parrafo2.font
    fuente_parrafo2.size = Pt(12)
    fuente_parrafo2.name = "Arial"
    fuente_parrafo3 = parrafo3.font
    fuente_parrafo3.size = Pt(12)
    fuente_parrafo3.name = "Arial"
    parrafo4.bold = True
    fuente_parrafo4 = parrafo4.font
    fuente_parrafo4.size = Pt(12)
    fuente_parrafo4.name = "Arial"
    fuente_parrafo5 = parrafo5.font
    fuente_parrafo5.size = Pt(12)
    fuente_parrafo5.name = "Arial"

    # Agregar la tabla con cada una de las activiades realizdas
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Fecha"
    hdr_cells[1].text = "Entrada"
    hdr_cells[2].text = "Salida"
    hdr_cells[3].text = "Total"
    hdr_cells[4].text = "Actividades Realizadas"
    hdr_cells[5].text = "Supervisor"

    for actividad in estudiante.actividades:
        row_cells = table.add_row().cells
        row_cells[0].text = actividad.fecha.strftime("%m/%d/%Y")
        row_cells[1].text = actividad.hora_inicio
        row_cells[2].text = actividad.hora_fin
        row_cells[3].text = str(actividad.cantidad_de_horas)
        row_cells[4].text = actividad.nombre_actividad
        row_cells[5].text = actividad.bibliotecario.nombre

    # Ciudad y Fecha
    ciudad_fechas = document.add_paragraph()
    ciudad_fechas.add_run("\n")
    ciudadyfecha = ciudad_fechas.add_run(f"Cuenca, {fecha}")
    fuente_ciudad_fecha = ciudadyfecha.font
    fuente_ciudad_fecha.size = Pt(12)
    fuente_ciudad_fecha.name = "Arial"
    formato_ciudad_fechas = ciudad_fechas.paragraph_format
    formato_ciudad_fechas.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    firma = document.add_paragraph()
    firma_coordinador = firma.add_run(f"\n\n\n{coordinador}")
    font_coordinador = firma_coordinador.font
    font_coordinador.size = Pt(10)
    font_coordinador.name = "Arial"
    formato_firma = firma.paragraph_format
    formato_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(certif_estud)
    return certif_estud
